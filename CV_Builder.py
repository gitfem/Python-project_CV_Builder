# importing Document from docx
from docx import Document
# From import inches, defining the size
from docx.shared import Inches

# assgining the name "document" to Document
document = Document()

# Adding profile Pictures
document.add_picture(
    'Emiju_Caleb_Femi passport.jpeg', 
    width=Inches(2.0)
)

# Asking for user input of the following:
document.add_heading('Personal Details')
name = input('What is Name ')
phone_number = input('What is Phone Number ')
email = input('What is Your Email Address ')

# Adding input colleted to the document as a paragraph
document.add_paragraph(
     'Name' + ':- '   + name + '\n' + 'Phone Number' + ':- '   + phone_number  + '\n' + 'Email Address' + ':- '   + email ).bold = True
      
# About me heading
document.add_heading('About me')
#about_me = input('Tell me about youself') 
#document.add_paragraph(about_me)
# OR
document.add_paragraph(input('Tell me about yourself '))

# Adding work experience heading and defind varable for p
document.add_heading('Work Experience')
p = document.add_paragraph()

# Asking for User input of the following:
company = input('Enter Company ')
from_date = input('From Date ')
to_date = input('To Date ')

# Formatting the paragraph of the user input
p.add_run(company + ' ').bold = True
p.add_run(from_date + '-' + to_date + '\n').italic = True

experience_details = input('Decribe your work experience at ' + company )
p.add_run(experience_details)

# adding more experiences "yes or no" using while loop
while True:
    has_more_experiences = input(
        'Do you have more experiences? Yes or No '
    )
    if has_more_experiences.lower() == 'yes':
        p = document.add_paragraph()

        # Asking for User input of the following:
        company = input('Enter Company ')
        from_date = input('From Date ')
        to_date = input('To Date ')

        # Formatting the paragraph of the user input
        p.add_run(company + ' ').bold = True
        p.add_run(from_date + '-' + to_date + '\n').italic = True

        experience_details = input('Decribe your work experience at ' + company )
        p.add_run(experience_details)
    else:
        break

# Adding Skils
document.add_heading('Skills')
skill = input('Add your Skill ')
s = document.add_paragraph(skill)
s.style = 'List Bullet'

while True:
    has_more_skills = input('Do you have more Skills? Yes or No ')
    if has_more_skills.lower() == 'yes':
        skills = input('Enter Skill ')
        s = document.add_paragraph(skill)
        s.style = 'List Bullet'  
    else:
        break

# Adding Footer
section = document.sections[0]
footer = section.footer
p = footer.paragraphs[0]
p.text = "CV generated using Techwizi and QuickBooks CV Builder project"

# Saving the document as the file name "CV_Builder.docx"

document.save('CV_Builder.docx')